# -*- coding: cp1252 -*-
#TO EXTRACT NAMES OF CHARACTERS(WHICH ARE IN BOLD) FROM .docx file
from docx import *
import time
import numpy as np
from scipy.stats import zscore
import itertools
import collections
plot = 1
start = time.time()
document = Document("D:\\Users\\Hrishi\\Documents\\PESU\\5th Semester\\Advanced Algorithms\\Project\\Ramayana.docx")
bolds=[]
for para in document.paragraphs:
    for run in para.runs:
        if run.bold :
            word=run.text
            if word!="":
                if word[0].isupper():
                    
                    bolds.append(word)
#print(*bolds,sep="\n")
    
sno=1
count=0  
character=dict()
unique=set(bolds)
#print(unique)   
for i in unique:
    for j in bolds:
        if i==j:
            count+=1
    if i!="":
        character[i]=count
    count=0;
    
ranked_list=sorted(character, key=lambda x: character[x])[::-1]
for i in ranked_list:
    #print(sno,")",i,character[i])
    sno+=1



#print(ranked_list[0:14])

#TO NAVIGATE WORDS AND SENTENCES IN .TXT FILE WHICH HAS THE SAME CONTENTS AS THE .DOCX FILE
f=open("D:\\Users\\Hrishi\\Documents\\PESU\\5th Semester\\Advanced Algorithms\\Project\\Ramayana.txt","r")
text=f.read()
sentences=text.split(".")
relation=[]#contains a list of tuples of characters appearing together
rel=[]#contains two characters that appear together
c=0
for sentence in sentences:
    words=sentence.split()
    for i in range(0,len(words)):
        #Creates list of two characters appearing togetheR        
        if words[i] in bolds:
            rel.append(words[i])
        else:
            continue
        for j in range(i+1,len(words)):
            if words[j] in bolds:
                if(len(rel)<2 and words[j]!=rel[0]):#Createst the final pair
                    rel.append(words[j])
                    relation.append(tuple(rel))
                    rel=[words[i]]
        rel=[]#Reinitializing for next pair
                
                
key_set=[]
relationship=dict()
for i in range(0, len(relation)):
    item=relation[i]
    count=1
    if(set(item) not in key_set):
        for j in range(i+1, len(relation)):
            if(set(item)==set(relation[j])):
                count+=1
        key_set.append(set(item))
        relationship[item]=count
        count=0
'''       
#RANKING IN ORDER OF IMPORTANCE OF RELATIONSHIP
ranked_list=sorted(relationship, key=lambda x: relationship[x])[::-1]
#print(ranked_list)
n = [i for i in unique]
for i in ranked_list:
    if(n_ in i):
        print(i,relationship[i])

'''
f.close()







document = Document("D:\\Users\\Hrishi\\Documents\\PESU\\5th Semester\\Advanced Algorithms\\Project\\Ramayana.docx")
bolds=[]
for para in document.paragraphs:
    for run in para.runs:
        if run.bold :
            word=run.text
            if word!="":
                if word[0].isupper():
                    
                    bolds.append(word)

count=0  
character=dict()
unique=set(bolds)
#print(unique)   
for i in unique:
    for j in bolds:
        if i==j:
            count+=1
    if i!="":
        character[i]=count
    count=0;
    
bolds=sorted(character, key=lambda x: character[x])[::-1]
#print(bolds)
male_pronouns=["he","his","He","His","him","Him","sons","twins","twin"]
female_pronouns=["she","hers","her","She","Hers","Her","mother","Mother","wife"]

f=open("D:\\Users\\Hrishi\\Documents\\PESU\\5th Semester\\Advanced Algorithms\\Project\\Ramayana.txt","r")
text=f.read()
sentences=text.split(".")
male=dict()
female=dict()
mcount=0
fcount=0
k=0
for i in bolds:
    #print(i)
    for j in sentences:
        words=j.split()
        for word in words:
            if word==i: 
                for gender in words:
                    if gender in male_pronouns:
                        #print("MALE")
                        mcount+=1
                    elif gender in female_pronouns:
                        #print("FEMALE")
                        fcount+=1
                            
        
    """if i=="Kusha":
        print(mcount)
        print(fcount)
    """
    male[i]=mcount
    female[i]=fcount
    mcount=0
    fcount=0

final=dict()
males=[]
females=[]
ambiguous=[]
for i in bolds:
    #print(i)
    if male[i]>female[i]:
        final[i]="M"
        males.append(i)
    elif female[i]>male[i]:
        final[i]="F"
        females.append(i)
    else:
        ambiguous.append(i)

monkeys = list()
monkeys = ['Vali','Hanuman','Sugriva','Kabandha']









graph = dict()
graph_transitive = dict()
for char in unique:
    graph[char]=dict()
    graph_transitive[char]=dict()
    for char_ in unique:
        graph[char][char_] = 0
        graph_transitive[char][char_] = 0
        
        
for c1,c2 in relation:
    try:
        graph[c1][c2]+=1
        graph[c2][c1]+=1
        graph_transitive[c1][c2] = 1
        graph_transitive[c2][c1] = 1
    except Exception as e:
        print(e)
        pass

graph_transitive_clousure = graph_transitive

print("Enter two Characters")
a = input("Enter char 1")
b = input("Enter char 2")



for k in (graph):
    for i in (graph):
        for j in (graph):
            try:
                graph_transitive_clousure[i][j] = graph_transitive_clousure[i][j] or (graph_transitive_clousure[i][k] and graph_transitive_clousure[k][j]) 
            except:
                print(k,i,j)


print("Strength of their relationship is")
print(graph[a][b])


new_graph = dict()
for x,y in relation:
    if x not in new_graph:
        new_graph[x]=list()
    if not y in new_graph[x]:
        new_graph[x].append(y)
for char in unique:
    if char not in new_graph:    
        new_graph[char] = list()


#print(new_graph)
for i in unique:
    graph_transitive_clousure[i][i]=0



if (plot):
    length = 1804 
    plotx = []
    for x in graph:
        for y in graph:
            plotx.append(graph[x][y])
    plotx.sort(reverse = True)
    ploty = list(range(1,len(plotx)+1))
    import math
    from matplotlib import pyplot as plt
    from scipy.stats import pearsonr
    plotxLog = list(map(lambda x: math.log(x),plotx[:-length]))
    plotyLog = list(map(lambda x: math.log(x),ploty[:-length]))
    plt.plot(ploty[:-length],plotx[:-length])
    plt.xlabel('Relationship Rank')
    plt.ylabel('Strength of relation')
    plt.show()
    plt.plot(plotyLog,plotxLog)
    plt.xlabel('Log(Relationship Rank)')
    plt.ylabel('Log(Strength of relation)')
    plt.show()
    corr, _ = pearsonr(plotxLog, plotyLog)
    print("Alpah",abs(corr)) 
    



recurssive_counter = 0
no_of_chars = 5
